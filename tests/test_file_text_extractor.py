"""file_text_extractor 单元测试（设计 docs/design/smart-import-file-extraction.md §11）。

fixture 在测试内用 openpyxl / python-docx 临时写到 tmp_path（不入库二进制样本，
与项目现有风格一致）。各格式覆盖：正常解析、空文档、公式回退、损坏抛错、回退纯文本、
source_format 标识。
"""

from __future__ import annotations

from pathlib import Path

import pytest

from management_prd.errors import LlmError
from management_prd.services.file_text_extractor import extract_text_for_llm

# ── Excel .xlsx ──


def _write_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "主界面"
    ws1.append(["模块", "功能", "状态"])
    ws1.append(["登录", "实现微信登录", "已完成"])
    ws1.append(["支付", "实现支付", "待办"])
    ws2 = wb.create_sheet("账户")
    ws2.append(["字段", "说明"])
    ws2.append(["手机号", "必填"])
    wb.save(path)


def test_xlsx_multi_sheet(tmp_path: Path) -> None:
    """多 sheet：输出含每个 `## 工作表: {name}`、Markdown 表格行列正确、首行作表头。"""
    path = tmp_path / "req.xlsx"
    _write_xlsx(path)

    text, fmt = extract_text_for_llm(path)

    assert fmt == "xlsx"
    assert "# 工作簿: req（共 2 个工作表）" in text
    assert "## 工作表: 主界面" in text
    assert "## 工作表: 账户" in text
    # 首行作表头 + 分隔行
    assert "| 模块 | 功能 | 状态 |" in text
    assert "| --- | --- | --- |" in text
    assert "| 登录 | 实现微信登录 | 已完成 |" in text


def test_xlsx_empty_sheet(tmp_path: Path) -> None:
    """空 sheet：输出（空工作表），不抛错。"""
    from openpyxl import Workbook

    path = tmp_path / "empty.xlsx"
    wb = Workbook()
    wb.active.title = "空表"  # 不写入任何数据
    wb.save(path)

    text, fmt = extract_text_for_llm(path)

    assert fmt == "xlsx"
    assert "## 工作表: 空表" in text
    assert "（空工作表）" in text


def test_xlsx_formula_fallback(tmp_path: Path) -> None:
    """data_only 缓存为 None 的公式单元格回退取公式串，非空（决策 4）。"""
    from openpyxl import Workbook

    path = tmp_path / "formula.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "结果"
    ws["A2"] = "=SUM(1,2)"  # openpyxl 不计算，data_only 读取为 None
    wb.save(path)

    text, _ = extract_text_for_llm(path)

    # 公式串回退后非空
    assert "=SUM(1,2)" in text


def test_xlsx_corrupt_raises(tmp_path: Path) -> None:
    """损坏的 xlsx（伪 zip）-> LlmError 含「无法读取 Excel」。"""
    path = tmp_path / "bad.xlsx"
    path.write_bytes(b"not a real xlsx zip payload")

    with pytest.raises(LlmError) as exc:
        extract_text_for_llm(path)
    assert "无法读取 Excel" in str(exc.value)


def test_xlsx_merged_cells_filled(tmp_path: Path) -> None:
    """合并单元格：被合并区单元格在 Markdown 中填充左上角值，不再留空（决策见模块 docstring）。"""
    from openpyxl import Workbook

    path = tmp_path / "merged.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["模块", "功能", "状态"])
    ws.append(["登录与支付", "微信登录", "已完成"])
    ws.append(["登录与支付", "找回密码", "待办"])
    ws.append(["登录与支付", "微信支付", "待办"])
    ws.merge_cells("A2:A4")
    wb.save(path)

    text, _ = extract_text_for_llm(path)

    # 合并区每行都出现「登录与支付」，不再有空列
    assert "| 登录与支付 | 微信登录 | 已完成 |" in text
    assert "| 登录与支付 | 找回密码 | 待办 |" in text
    assert "| 登录与支付 | 微信支付 | 待办 |" in text
    # 横向合并（B2:C2 跨列）也填充
    ws2 = Workbook()
    w2 = ws2.active
    w2.append(["模块", "说明"])
    w2.append(["登录", "包括注册与找回密码的内容很长"])
    w2.merge_cells("B2:C2")
    w2["B2"] = "登录模块全部功能"
    path2 = tmp_path / "merged_h.xlsx"
    ws2.save(path2)
    text2, _ = extract_text_for_llm(path2)
    assert "| 登录 | 登录模块全部功能 |" in text2


def test_xlsx_dates_named(tmp_path: Path) -> None:
    """日期归一：datetime 对象、日期序列号+中文格式、中文/斜杠/横线/点日期串统一转 ISO。"""
    from datetime import datetime

    from openpyxl import Workbook

    path = tmp_path / "dates.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["名称", "日期"])
    ws.append(["datetime对象", datetime(2024, 11, 12)])
    ws.append(["序列号+中文格式", 45608.0])
    ws.append(["序列号+日期格式", 45608.0])
    ws.append(["中文文本", "2024年11月12日"])
    ws.append(["斜杠文本", "2024/11/13"])
    ws.append(["横线文本", "2024-12-01"])
    ws.append(["点文本", "2024.12.05"])
    # 序列号+日期格式：给 B 列对应行设日期格式触发序列号转日期分支
    ws["B3"].number_format = 'yyyy"年"m"月"d"日"'
    ws["B4"].number_format = "yyyy-mm-dd"
    wb.save(path)

    text, _ = extract_text_for_llm(path)

    assert "| datetime对象 | 2024-11-12 |" in text
    assert "| 序列号+中文格式 | 2024-11-12 |" in text
    assert "| 序列号+日期格式 | 2024-11-12 |" in text
    assert "| 中文文本 | 2024-11-12 |" in text
    assert "| 斜杠文本 | 2024-11-13 |" in text
    assert "| 横线文本 | 2024-12-01 |" in text
    assert "| 点文本 | 2024-12-05 |" in text


def test_xlsx_plain_number_not_converted(tmp_path: Path) -> None:
    """非日期格式的普通数值不应被误转成日期（45608 若格式为 General 保持原样）。"""
    from openpyxl import Workbook

    path = tmp_path / "num.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["名称", "数量"])
    ws.append(["订单数", 45608])  # General 格式，非日期
    wb.save(path)

    text, _ = extract_text_for_llm(path)

    assert "| 订单数 | 45608 |" in text


def test_xlsx_description_with_date_not_converted(tmp_path: Path) -> None:
    """描述性文本里的日期不应被整段替换：仅整串为日期才归一。"""
    from openpyxl import Workbook

    path = tmp_path / "desc.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["模块", "说明"])
    ws.append(["登录", "预计2024年11月12日上线，含微信支付"])  # 日期嵌在描述中
    wb.save(path)

    text, _ = extract_text_for_llm(path)

    assert "| 登录 | 预计2024年11月12日上线，含微信支付 |" in text


def test_xlsx_amplifies_text(tmp_path: Path) -> None:
    """Markdown 分隔符放大文本：抽取后长度 > 所有单元格值字符总和（长度校验须在抽取后）。"""
    from openpyxl import Workbook

    path = tmp_path / "big.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["模块", "功能", "状态"])
    values_sum = 0
    for i in range(50):
        row = [f"模块{i}", f"功能{i}", "已完成"]
        ws.append(row)
        values_sum += sum(len(v) for v in row)
    wb.save(path)

    text, _ = extract_text_for_llm(path)
    # 表格分隔符 / 表头分隔行带来额外字符 -> 抽取放大
    assert len(text) > values_sum


# ── Word .docx ──


def _write_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("标题段落")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "模块"
    table.cell(0, 1).text = "功能"
    table.cell(1, 0).text = "登录"
    table.cell(1, 1).text = "微信登录"
    doc.add_paragraph("结尾段落")
    doc.save(path)


def test_docx_paragraph_table_order(tmp_path: Path) -> None:
    """段落+表格交错：输出保序（段落在前则段落在前），表格转 Markdown。"""
    path = tmp_path / "doc.docx"
    _write_docx(path)

    text, fmt = extract_text_for_llm(path)

    assert fmt == "docx"
    assert "# 文档: doc" in text
    assert "标题段落" in text
    assert "结尾段落" in text
    assert "| 模块 | 功能 |" in text
    assert "| --- | --- |" in text
    assert "| 登录 | 微信登录 |" in text
    # 保序：标题段落 在 表格 之前，表格 在 结尾段落 之前
    assert text.find("标题段落") < text.find("| 模块 | 功能 |")
    assert text.find("| 模块 | 功能 |") < text.find("结尾段落")


def test_docx_empty(tmp_path: Path) -> None:
    """空文档：仅标题行，不抛错。"""
    from docx import Document

    path = tmp_path / "empty.docx"
    Document().save(path)  # 默认模板仅含空段落（被跳过）

    text, fmt = extract_text_for_llm(path)

    assert fmt == "docx"
    assert text.strip() == "# 文档: empty"


# ── .xls / .doc / 纯文本 / 未知扩展名 ──


def _write_xls(path: Path) -> None:
    """用 xlwt 生成最小 .xls fixture（xlrd 只读，写需 xlwt）。"""
    import xlwt

    wb = xlwt.Workbook()
    ws1 = wb.add_sheet("主界面")
    ws1.write(0, 0, "模块")
    ws1.write(0, 1, "功能")
    ws1.write(1, 0, "登录")
    ws1.write(1, 1, "微信登录")
    ws2 = wb.add_sheet("账户")
    ws2.write(0, 0, "字段")
    ws2.write(0, 1, "说明")
    wb.save(str(path))


def test_xls_multi_sheet(tmp_path: Path) -> None:
    """.xls 多 sheet：输出 `## 工作表` 标题 + Markdown 表格，source_format=xls。"""
    path = tmp_path / "req.xls"
    _write_xls(path)

    text, fmt = extract_text_for_llm(path)

    assert fmt == "xls"
    assert "# 工作簿: req（共 2 个工作表）" in text
    assert "## 工作表: 主界面" in text
    assert "## 工作表: 账户" in text
    assert "| 模块 | 功能 |" in text
    assert "| --- | --- |" in text
    assert "| 登录 | 微信登录 |" in text


def test_xls_corrupt_raises(tmp_path: Path) -> None:
    """损坏 .xls -> LlmError 含「无法读取 Excel」。"""
    path = tmp_path / "bad.xls"
    path.write_bytes(b"not a real xls")

    with pytest.raises(LlmError) as exc:
        extract_text_for_llm(path)
    assert "无法读取 Excel" in str(exc.value)


def test_doc_not_supported(tmp_path: Path) -> None:
    """.doc -> LlmError 含「另存为 .docx」。"""
    path = tmp_path / "old.doc"
    path.write_bytes(b"anything")

    with pytest.raises(LlmError) as exc:
        extract_text_for_llm(path)
    assert "另存为 .docx" in str(exc.value)


@pytest.mark.parametrize(
    "name,expected_fmt",
    [
        ("doc.txt", "txt"),
        ("notes.md", "md"),
        ("data.csv", "csv"),
    ],
)
def test_text_formats(tmp_path: Path, name: str, expected_fmt: str) -> None:
    """文本类扩展名：与 read_text(errors='replace') 等价，source_format 正确。"""
    path = tmp_path / name
    content = "登录：实现微信登录\n支付：实现支付"
    path.write_text(content, encoding="utf-8")

    text, fmt = extract_text_for_llm(path)

    assert fmt == expected_fmt
    assert text == content


def test_unknown_ext_falls_back_to_text(tmp_path: Path) -> None:
    """未知扩展名 -> 回退纯文本读取（现状），source_format 取扩展名。"""
    path = tmp_path / "weird.log2"
    path.write_text("一些内容", encoding="utf-8")

    text, fmt = extract_text_for_llm(path)

    assert fmt == "log2"
    assert text == "一些内容"
